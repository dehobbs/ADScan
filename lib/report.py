""" lib/report.py - ADScan HTML Report Generator
Generates a self-contained HTML dashboard with:
  - Light/Dark mode toggle
  - Security score gauge
  - LEFT SIDEBAR: category navigation (alphabetical, collapsible, ordered by severity)
  - Clickable, multi-select severity filter chips
  - Per-finding cards with severity badges
  - Scan metadata table
  - Manual Verification tool cards (per finding)
  - Remediation guidance (per finding)
  - Executive summary with score, grade, and top priority findings
"""
from datetime import datetime
import html as html_mod
import importlib
import os
import pkgutil

# ---------------------------------------------------------------------------
# Severity colour mapping
# ---------------------------------------------------------------------------
SEVERITY_COLORS = {
    "critical": ("#dc2626", "#fef2f2", "#7f1d1d"),
    "high":     ("#ea580c", "#fff7ed", "#7c2d12"),
    "medium":   ("#d97706", "#fffbeb", "#78350f"),
    "low":      ("#2563eb", "#eff6ff", "#1e3a8a"),
    "info":     ("#6b7280", "#f9fafb", "#374151"),
}

SEV_ORDER = ["critical", "high", "medium", "low", "info"]
_SEV_RANK = {s: i for i, s in enumerate(SEV_ORDER)}


def _score_color(score):
    if score >= 90: return "#16a34a"
    if score >= 75: return "#84cc16"
    if score >= 60: return "#eab308"
    if score >= 40: return "#f97316"
    return "#dc2626"


def _grade(score):
    if score >= 90: return "A"
    if score >= 75: return "B"
    if score >= 60: return "C"
    if score >= 40: return "D"
    return "F"


def _references_html(finding):
    """Render the References section for a finding."""
    vdata = _get_verification(finding)
    if not vdata or not vdata.get("references"):
        return ""
    refs = vdata["references"]
    items_html = ""
    for ref in refs:
        label = html_mod.escape(ref.get("title", "Reference"))
        url = ref.get("url", "#")
        tag = ref.get("tag", "")
        tag_html = ""
        if tag:
            tag_colors = {
                "vendor": ("#1d4ed8", "#dbeafe"),
                "attack": ("#991b1b", "#fee2e2"),
                "defense": ("#065f46", "#d1fae5"),
                "research": ("#5b21b6", "#ede9fe"),
                "tool": ("#92400e", "#fef3c7"),
            }
            tc, tbc = tag_colors.get(tag.lower(), ("#374151", "#f3f4f6"))
            tag_html = (
                f'<span style="font-size:0.68rem;font-weight:700;'
                f'text-transform:uppercase;letter-spacing:0.06em;'
                f'background:{tbc};color:{tc};border-radius:4px;'
                f'padding:1px 6px;margin-right:6px;flex-shrink:0;">'
                f'{html_mod.escape(tag)}</span>'
            )
        items_html += (
            f'<li style="display:flex;align-items:baseline;gap:4px;'
            f'margin-bottom:6px;font-size:0.85rem;">'
            f'{tag_html}'
            f'<a href="{html_mod.escape(url)}" target="_blank" rel="noopener noreferrer" '
            f'style="color:var(--ref-link);text-decoration:none;word-break:break-all;">'
            f'{label}</a>'
            f'</li>'
        )
    ref_count = len(refs)
    return f"""<details style='margin-top:12px;'>
  <summary style='cursor:pointer;font-weight:600;color:var(--text-muted);
    font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;
    list-style:none;display:flex;align-items:center;gap:6px;user-select:none;'>
    <span style='display:inline-block;transition:transform .2s;font-size:0.65rem;'>&#9660;</span>
    References ({ref_count})
  </summary>
  <div style='margin-top:10px;padding:12px 16px;background:var(--ref-bg);
    border:1px solid var(--ref-border);border-radius:8px;'>
    <ul style='list-style:none;margin:0;padding:0;'>
      {items_html}
    </ul>
  </div>
</details>"""


def _severity_badge_html(severity):
    sev = severity.lower()
    bg_light, _, _ = SEVERITY_COLORS.get(sev, ("#6b7280", "#f9fafb", "#374151"))
    return (
        f'<span class="badge badge-{sev}" '
        f'style="background:{bg_light};color:#fff;'
        f'padding:2px 10px;border-radius:12px;font-size:0.78rem;'
        f'font-weight:600;letter-spacing:0.05em;text-transform:uppercase;">'
        f'{html_mod.escape(severity.upper())}</span>'
    )


# ---------------------------------------------------------------------------
# Manual Verification & Remediation Database
# ---------------------------------------------------------------------------
# Keys are lowercase substrings matched against finding["title"].lower().
# First match wins. Tools list drives the 2-column verification grid.
# ---------------------------------------------------------------------------


def _build_verification_db():
    """Auto-discover verification modules from the verifications/ package."""
    import sys as _sys
    db = {}
    verif_path = os.path.join(os.path.dirname(__file__), '..', 'verifications')
    for _finder, name, _ispkg in pkgutil.iter_modules([verif_path]):
        try:
            mod = importlib.import_module(f'verifications.{name}')
        except Exception as _e:
            print(f'[ADScan] Warning: could not load verifications/{name}.py: {_e}', file=_sys.stderr)
            continue
        if hasattr(mod, 'TOOLS'):
            entry = {
                'tools': mod.TOOLS,
                'remediation': getattr(mod, 'REMEDIATION', None),
                'references': getattr(mod, 'REFERENCES', None),
            }
            match_keys = getattr(mod, 'MATCH_KEYS', [])
            for key in match_keys:
                db[key] = entry
    return db


VERIFICATION_DB = _build_verification_db()


def _get_verification(finding):
    """Return the VERIFICATION_DB entry that matches this finding, or None."""
    title = finding.get("title", "").lower()
    for key, data in VERIFICATION_DB.items():
        if key in title:
            return data
    return None


def _tool_icon_html(icon_type):
    """Return a small coloured SVG/text icon square for a tool card header."""
    icons = {
        "netexec": (
            "#1a1a2e",
            '<svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" viewBox="0 0 24 24" fill="white">'
            '<rect x="2" y="2" width="4" height="4" rx="1"/><rect x="10" y="2" width="4" height="4" rx="1"/>'
            '<rect x="18" y="2" width="4" height="4" rx="1"/><rect x="2" y="10" width="4" height="4" rx="1"/>'
            '<rect x="10" y="10" width="4" height="4" rx="1"/><rect x="18" y="10" width="4" height="4" rx="1"/>'
            '<rect x="2" y="18" width="4" height="4" rx="1"/><rect x="10" y="18" width="4" height="4" rx="1"/>'
            '<rect x="18" y="18" width="4" height="4" rx="1"/></svg>'
        ),
        "impacket": (
            "#2d3748",
            '<svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" viewBox="0 0 24 24" fill="white">'
            '<path d="M12 2L2 7l10 5 10-5-10-5zM2 17l10 5 10-5M2 12l10 5 10-5"/></svg>'
        ),
        "ps": (
            "#012456",
            '<svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" viewBox="0 0 24 24" fill="white">'
            '<text x="2" y="17" font-family="monospace" font-size="12" font-weight="bold">PS</text></svg>'
        ),
        "cmd": (
            "#1a1a1a",
            '<svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" viewBox="0 0 24 24" fill="white">'
            '<polyline points="4 9 9 12 4 15" stroke="white" stroke-width="2" fill="none" stroke-linecap="round"/>'
            '<line x1="12" y1="15" x2="20" y2="15" stroke="white" stroke-width="2" stroke-linecap="round"/></svg>'
        ),
        "aduc": (
            "#c9943a",
            '<svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" viewBox="0 0 24 24" fill="white">'
            '<text x="3" y="17" font-family="serif" font-size="15" font-weight="bold">A</text></svg>'
        ),
    }
    bg, svg = icons.get(icon_type, icons["cmd"])
    return (
        f'<span class="verif-icon" style="background:{bg};">{svg}</span>'
    )


def _tool_card_html(tool_data):
    """Render one tool verification card."""
    icon_html = _tool_icon_html(tool_data.get("icon", "cmd"))
    tool_name = html_mod.escape(tool_data.get("tool", "Tool"))

    body_parts = []

    # Description
    desc = tool_data.get("desc", "")
    if desc:
        body_parts.append(f'<p class="verif-desc">{desc}</p>')

    # Numbered steps (for ADUC-style cards)
    steps = tool_data.get("steps", [])
    if steps:
        steps_html = "".join(
            f'<li>{s}</li>' for s in steps
        )
        body_parts.append(f'<ol class="verif-steps">{steps_html}</ol>')

    # Code block
    code = tool_data.get("code", "")
    if code:
        body_parts.append(
            f'<pre class="verif-code"><code>{html_mod.escape(code)}</code></pre>'
        )

    # Confirmation text
    confirm = tool_data.get("confirm", "")
    if confirm:
        body_parts.append(f'<p class="verif-confirm"><em>{confirm}</em></p>')

    body_html = "\n".join(body_parts)

    return f"""<div class="verif-card">
  <div class="verif-card-header">
    {icon_html}
    <span class="verif-tool-name">{tool_name}</span>
  </div>
  {body_html}
</div>"""


def _manual_verification_html(finding):
    """Render the Manual Verification section as Linux/Windows tabs, tiles stacked."""
    vdata = _get_verification(finding)
    if not vdata or not vdata.get("tools"):
        return ""

    tools = vdata["tools"]

    # Split tools by platform based on icon type
    _LINUX_ICONS = {"netexec", "impacket"}
    linux_tools = [t for t in tools if t.get("icon", "cmd") in _LINUX_ICONS]
    win_tools   = [t for t in tools if t.get("icon", "cmd") not in _LINUX_ICONS]

    # Each tab panel: stacked single-column list of tool cards
    def _panel(tool_list):
        if not tool_list:
            return ""
        return "\n".join(_tool_card_html(t) for t in tool_list)

    # Build unique tab-group ID so multiple findings on the page don't clash
    import hashlib
    tab_id = "vt-" + hashlib.md5(str(id(finding)).encode(), usedforsecurity=False).hexdigest()[:8]

    tabs_html = ""
    panels_html = ""
    first = True

    if linux_tools:
        active_tab = "verif-tab-active" if first else "verif-tab-inactive"
        active_panel = "" if first else " style='display:none;'"
        lx_svg = '<svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2.2" style="flex-shrink:0"><rect x="2" y="3" width="20" height="14" rx="2"/><line x1="8" y1="21" x2="16" y2="21"/><line x1="12" y1="17" x2="12" y2="21"/></svg>'
        tabs_html += f'<button class="verif-tab {active_tab}" onclick="verifTab(this,\'{tab_id}-linux\')">{lx_svg}Linux</button>'
        panels_html += f'<div id="{tab_id}-linux" class="verif-panel"{active_panel}>{_panel(linux_tools)}</div>'
        first = False

    if win_tools:
        active_tab = "verif-tab-active" if first else "verif-tab-inactive"
        active_panel = "" if first else " style='display:none;'"
        win_svg = '<svg width="13" height="13" viewBox="0 0 24 24" fill="currentColor" style="flex-shrink:0"><path d="M0 3.5L10 2v10H0V3.5zm0 17L10 22V12H0v8.5zM11 1.8L24 0v12H11V1.8zm0 20.2L24 24V12H11v10z"/></svg>'
        tabs_html += f'<button class="verif-tab {active_tab}" onclick="verifTab(this,\'{tab_id}-win\')">{win_svg}Windows</button>'
        panels_html += f'<div id="{tab_id}-win" class="verif-panel"{active_panel}>{_panel(win_tools)}</div>'

    card_count = len(tools)
    return f"""<details style='margin-top:16px;'>
  <summary style='cursor:pointer;font-weight:600;color:var(--text-muted);
    font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;
    list-style:none;display:flex;align-items:center;gap:6px;user-select:none;'>
    <span style='display:inline-block;transition:transform .2s;font-size:0.65rem;'>&#9660;</span>
    Manual Verification ({card_count} tool{"s" if card_count != 1 else ""})
  </summary>
  <div class="verif-tabs-wrap">
    <div class="verif-tab-bar">{tabs_html}</div>
    <div class="verif-panels">{panels_html}</div>
  </div>
</details>"""
def _remediation_html(finding):
    """Render the Remediation section for a finding."""
    vdata = _get_verification(finding)
    if not vdata or not vdata.get("remediation"):
        return ""

    remed = vdata["remediation"]
    title = html_mod.escape(remed.get("title", "Remediation"))
    steps = remed.get("steps", [])

    steps_html = ""
    for i, step in enumerate(steps, 1):
        step_text = step.get("text", "")
        code = step.get("code", "")
        sub_steps = step.get("steps", [])
        code_block = ""
        if code:
            code_block = f'<pre class="verif-code remed-code"><code>{html_mod.escape(code)}</code></pre>'
        sub_steps_html = ""
        if sub_steps:
            items = "".join(f"<li>{s}</li>" for s in sub_steps)
            sub_steps_html = f'<ol class="verif-steps" style="margin-top:6px;">{items}</ol>'
        steps_html += f"""<div class="remed-step">
        <span class="remed-num">{i}</span>
        <div class="remed-step-body">
            <p>{step_text}</p>
            {sub_steps_html}
            {code_block}
        </div>
        </div>"""

    step_count = len(steps)
    return f"""<details style='margin-top:12px;'>
  <summary style='cursor:pointer;font-weight:600;color:var(--text-muted);
    font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;
    list-style:none;display:flex;align-items:center;gap:6px;user-select:none;'>
    <span style='display:inline-block;transition:transform .2s;font-size:0.65rem;'>&#9660;</span>
    Remediation ({step_count} step{"s" if step_count != 1 else ""})
  </summary>
  <div class="remed-box" style='margin-top:12px;'>
    <div class="remed-title">Recommended: {title}</div>
    {steps_html}
  </div>
</details>"""


def _severity_badge_html(severity):
    sev = severity.lower()
    bg_light, _, _ = SEVERITY_COLORS.get(sev, ("#6b7280", "#f9fafb", "#374151"))
    return (
        f'<span class="badge badge-{sev}" '
        f'style="background:{bg_light};color:#fff;'
        f'padding:2px 10px;border-radius:12px;font-size:0.78rem;'
        f'font-weight:600;letter-spacing:0.05em;text-transform:uppercase;">'
        f'{html_mod.escape(severity.upper())}</span>'
    )


   



def _finding_card(finding, idx):
    severity = finding.get("severity", "info").lower()
    sev_colors = SEVERITY_COLORS.get(severity, SEVERITY_COLORS["info"])
    accent = sev_colors[0]
    category = finding.get("category", "Uncategorized")
    if isinstance(category, str):
        cat_list = [category]
    else:
        cat_list = list(category)

    details_html = ""
    if finding.get("details"):
        items = "".join(
            f"<li style='margin:2px 0;font-family:monospace;font-size:0.85rem;'>"
            f"{html_mod.escape(str(d)).replace('[[REDACTED]]', '<span style=\"color:#e53e3e;font-weight:bold\">REDACTED</span>')}</li>"
            for d in finding["details"][:50]
        )
        more = ""
        if len(finding["details"]) > 50:
            more = f"<li><em>... and {len(finding['details']) - 50} more</em></li>"
        details_html = f"""
    <details style='margin-top:12px;'>
      <summary style='cursor:pointer;font-weight:600;color:var(--text-muted);
        font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;
        list-style:none;display:flex;align-items:center;gap:6px;user-select:none;'>
        <span style='display:inline-block;transition:transform .2s;font-size:0.65rem;color:{accent}'>&#9660;</span>
        <span style='color:{accent}'>Affected Objects ({finding.get('affected_count', len(finding['details']))})</span>
      </summary>
      <ul style='margin:10px 0 0 16px;padding:0;'>{items}{more}</ul>
    </details>"""

    verif_html = _manual_verification_html(finding)
    remed_html = _remediation_html(finding)
    refs_html = _references_html(finding)

    cat_slug = cat_list[0].lower().replace(" ", "-").replace("&", "and").replace("/", "-")
    return f"""
<div class="finding-card" id="finding-{idx}"
     data-severity="{severity}"
     data-category="{" ".join(c.lower().replace(" ", "-").replace("&", "and").replace("/", "-") for c in cat_list)}"
     style="border-left:4px solid {accent};background:var(--card-bg);
            padding:20px 24px;margin-bottom:16px;border-radius:8px;
            box-shadow:var(--card-shadow);">
  <div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:8px;">
    <div>
      <h3 style="margin:0 0 6px;font-size:1.05rem;">{html_mod.escape(finding.get('title',''))}</h3>
      {_severity_badge_html(severity)}
      <span style="margin-left:8px;font-size:0.85rem;color:var(--text-muted);">
        Risk Deduction: <strong style="color:{accent};">-{finding.get('deduction', 0)} pts</strong>
      </span>
      <span style="margin-left:8px;font-size:0.78rem;color:var(--text-muted);
                   background:var(--rec-bg);border-radius:10px;padding:2px 8px;">
        {html_mod.escape(", ".join(cat_list))}
      </span>
    </div>
  </div>
  <p style="margin:14px 0 8px;line-height:1.6;color:var(--text-secondary);">
    {html_mod.escape(finding.get('description', ''))}
  </p>
  <div style="background:var(--rec-bg);border-radius:6px;padding:10px 14px;margin-top:8px;">
    <strong style="font-size:0.85rem;color:var(--text-muted);text-transform:uppercase;letter-spacing:0.05em;">
      Recommendation
    </strong>
    <p style="margin:4px 0 0;font-size:0.92rem;line-height:1.5;">
      {html_mod.escape(finding.get('recommendation', 'Review and remediate this finding.'))}
    </p>
  </div>
  {details_html}
  {verif_html}
  {remed_html}
  {refs_html}
</div>"""


def _exec_summary_html(domain, dc_host, scan_time, score, grade, score_color, sev_counts, findings):
    """Build the Executive Summary card HTML inserted at the top of each report."""
    # ---- prose paragraph ----
    total = len(findings)
    sev_parts = []
    for sev in SEV_ORDER:
        count = sev_counts.get(sev, 0)
        if count == 0:
            continue
        color = SEVERITY_COLORS[sev][0]
        label = sev.capitalize()
        sev_parts.append(
            f'<strong style="color:{color};">{count} {label}</strong>'
        )
    if len(sev_parts) > 1:
        sev_str = ", ".join(sev_parts[:-1]) + f", and {sev_parts[-1]}"
    elif sev_parts:
        sev_str = sev_parts[0]
    else:
        sev_str = "none"

    if total == 0:
        posture = "an excellent"
        issue_note = "indicating an excellent security posture with no issues detected."
    elif score >= 90:
        posture = "an excellent"
        issue_note = (
            f"indicating {posture} security posture. "
            f"<strong>{total} informational finding{'s' if total != 1 else ''}</strong> "
            f"{'were' if total != 1 else 'was'} identified: {sev_str}."
        )
    elif score >= 75:
        issue_note = (
            f"indicating a good security posture with some issues requiring attention. "
            f"<strong>{total} finding{'s' if total != 1 else ''}</strong> "
            f"{'were' if total != 1 else 'was'} identified: {sev_str}."
        )
    elif score >= 60:
        issue_note = (
            f"indicating a moderate security posture with several issues requiring prompt attention. "
            f"<strong>{total} finding{'s' if total != 1 else ''}</strong> "
            f"{'were' if total != 1 else 'was'} identified: {sev_str}."
        )
    else:
        issue_note = (
            f"indicating a poor security posture with significant issues requiring immediate attention. "
            f"<strong>{total} finding{'s' if total != 1 else ''}</strong> "
            f"{'were' if total != 1 else 'was'} identified: {sev_str}."
        )

    urgency = ""
    if sev_counts.get("critical", 0) > 0 or sev_counts.get("high", 0) > 0:
        urgency = (
            " Immediate remediation of the critical and high-severity findings "
            "is strongly recommended before the next assessment cycle."
        )

    paragraph = (
        f'An authenticated security scan of <strong>{html_mod.escape(domain)}</strong> was '
        f'conducted on <strong>{html_mod.escape(scan_time[:10])}</strong> against domain '
        f'controller <strong>{html_mod.escape(dc_host)}</strong>. '
        f'The domain received a security score of '
        f'<strong style="color:{score_color};">{score}/100 (Grade {grade})</strong>, '
        f'{issue_note}{urgency}'
    )

    # ---- top-3 priority finding chips (findings already sorted by severity) ----
    top3 = findings[:3]
    chips_html = ""
    chip_bg_map = {
        "critical": ("#fef2f2", "#fecaca"),
        "high":     ("#fff7ed", "#fed7aa"),
        "medium":   ("#fffbeb", "#fde68a"),
        "low":      ("#eff6ff", "#bfdbfe"),
        "info":     ("#f9fafb", "#e5e7eb"),
    }
    for i, f in enumerate(top3):
        sev = f.get("severity", "info").lower()
        accent = SEVERITY_COLORS.get(sev, SEVERITY_COLORS["info"])[0]
        bg, border = chip_bg_map.get(sev, ("#f9fafb", "#e5e7eb"))
        cats = f.get("category", "Uncategorized")
        if not isinstance(cats, str):
            cats = ", ".join(cats)
        deduction = f.get("deduction", 0)
        chips_html += (
            f'<a class="exec-chip" href="#finding-{i}" style="background:{bg};border:1px solid {border};">'
            f'<div class="exec-chip-bar" style="background:{accent};"></div>'
            f'<div class="exec-chip-body">'
            f'<div class="exec-chip-sev" style="color:{accent};">{html_mod.escape(sev.upper())}</div>'
            f'<div class="exec-chip-title">{html_mod.escape(f.get("title", ""))}</div>'
            f'<div class="exec-chip-meta">-{deduction} pts &nbsp;&bull;&nbsp; {html_mod.escape(cats)}</div>'
            f'</div>'
            f'</a>'
        )

    top_section = ""
    if top3:
        n = len(top3)
        top_section = (
            f'<hr class="exec-divider">'
            f'<div class="exec-top-label">Top {n} Priority Finding{"s" if n != 1 else ""}</div>'
            f'<div class="exec-chips-row">{chips_html}</div>'
        )

    # ---- grade pill colours ----
    if score >= 75:
        pill_border, pill_bg = "#86efac", "#f0fdf4"
    elif score >= 50:
        pill_border, pill_bg = "#fde68a", "#fffbeb"
    else:
        pill_border, pill_bg = "#fca5a5", "#fef2f2"

    grade_pill = (
        f'<div class="exec-grade-pill" style="border-color:{pill_border};background:{pill_bg};">'
        f'<div class="exec-grade-circle" style="background:{score_color};">{grade}</div>'
        f'<div>'
        f'<div class="exec-grade-score" style="color:{score_color};">{score} / 100</div>'
        f'<div class="exec-grade-label">Security Score</div>'
        f'</div>'
        f'</div>'
    )

    icon_svg = (
        '<svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" '
        'stroke-width="2.2" style="flex-shrink:0;opacity:.55;">'
        '<path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/>'
        '<polyline points="14 2 14 8 20 8"/>'
        '<line x1="16" y1="13" x2="8" y2="13"/>'
        '<line x1="16" y1="17" x2="8" y2="17"/>'
        '</svg>'
    )

    return (
        f'<div class="exec-summary">'
        f'<div class="exec-header">'
        f'<div class="exec-title">{icon_svg} Executive Summary</div>'
        f'{grade_pill}'
        f'</div>'
        f'<p class="exec-body">{paragraph}</p>'
        f'{top_section}'
        f'</div>'
    )



def _category_scores_html(category_scores):
    """Build the category sub-scores grid HTML (Option 3)."""
    if not category_scores:
        return ""
    items_html = ""
    for cat, info in sorted(category_scores.items()):
        earned  = info.get("earned", 0)
        maximum = info.get("possible", 0)
        if maximum == 0:
            pct = 100
        else:
            pct = round(earned / maximum * 100)
        grade_str = _grade(pct)
        bar_color = _score_color(pct)
        cat_slug = cat.lower().replace(" ", "-").replace("&", "and").replace("/", "-")
        items_html += f"""<div class="cat-card" data-catslug="{cat_slug}" onclick="toggleCategoryFilter(this)" style="cursor:pointer;">
  <div class="cat-card-header">
    <span class="cat-card-name">{html_mod.escape(cat)}</span>
    <span class="cat-card-grade" style="background:{bar_color};">{grade_str}</span>
  </div>
  <div class="cat-bar-track">
    <div class="cat-bar-fill" style="width:{pct}%;background:{bar_color};"></div>
  </div>
  <div class="cat-card-score" style="color:{bar_color};">{earned}/{maximum} pts</div>
</div>"""
    return f"""<div class="cat-grid-section">
  <div class="cat-grid-label">Category Scores</div>
  <div class="cat-grid">{items_html}</div>
</div>"""


def generate_report(output_file, domain, dc_host, username, protocols, findings, score, category_scores=None):
    """Generate a self-contained HTML report dashboard with sidebar navigation."""
    from collections import defaultdict
    scan_time = __import__('datetime').datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    score_color = _score_color(score)
    grade = _grade(score)
    radius = 80
    circumference = 2 * 3.14159 * radius
    dash_offset = circumference * (1 - score / 100)

    # Severity counts
    sev_counts = {s: 0 for s in SEV_ORDER}
    for f in findings:
        sev = f.get("severity", "info").lower()
        if sev in sev_counts:
            sev_counts[sev] += 1

    # Sort findings: critical -> high -> medium -> low -> info
    findings = sorted(
        findings,
        key=lambda f: (
            _SEV_RANK.get(f.get("severity", "info").lower(), 99),
            -f.get("deduction", 0),
        ),
    )

    # ---- Sidebar ----
    cat_findings = defaultdict(list)
    for i, f in enumerate(findings):
        cats = f.get("category", "Uncategorized")
        if isinstance(cats, str):
            cats = [cats]
        for cat in cats:
            cat_findings[cat].append((i, f))

    sidebar_items_html = ""
    for cat in sorted(cat_findings.keys()):
        cat_id = cat.lower().replace(" ", "-").replace("&", "").replace("/", "-").replace(".", "")
        item_links = ""
        for idx, f in cat_findings[cat]:
            sev = f.get("severity", "info").lower()
            sev_color = SEVERITY_COLORS.get(sev, SEVERITY_COLORS["info"])[0]
            title_escaped = html_mod.escape(f.get("title", ""))
            item_links += (
                f'<a href="#finding-{idx}" class="sidebar-item" data-severity="{sev}"'
                f' onclick="sidebarNav(this,event)" title="{title_escaped}">'
                f'<span class="sev-dot" style="background:{sev_color};"></span>'
                f'<span class="sidebar-item-text">{title_escaped}</span>'
                f'</a>\n'
            )
        count = len(cat_findings[cat])
        sidebar_items_html += (
            f'<div class="cat-group" id="cat-{cat_id}">'
            f'<button class="cat-toggle" onclick="toggleCat(this)" aria-expanded="false">'
            f'<span class="cat-arrow">&#9660;</span>'
            f'<span class="cat-name">{html_mod.escape(cat)}</span>'
            f'<span class="cat-count">{count}</span>'
            f'</button>'
            f'<div class="cat-items collapsed">{item_links}</div>'
            f'</div>\n'
        )

    if not sidebar_items_html:
        sidebar_items_html = '<p class="sidebar-empty">No findings to navigate.</p>'

    # ---- Finding cards ----
    if findings:
        cards_html = "".join(_finding_card(f, i) for i, f in enumerate(findings))
    else:
        cards_html = (
            '<div style="text-align:center;padding:60px;color:var(--text-muted);">'
            '<div style="font-size:3rem;margin-bottom:16px;">&#10003;</div>'
            '<h3>No Vulnerabilities Found</h3>'
            '<p>All checks passed successfully.</p>'
            '</div>'
        )

    # ---- Severity chips ----
    severity_chips = ""
    for sev in SEV_ORDER:
        count = sev_counts.get(sev, 0)
        if count > 0:
            color = SEVERITY_COLORS[sev][0]
            cap_sev = sev.capitalize()
            severity_chips += (
                f'<button class="sev-chip" data-sev="{sev}" data-color="{color}"'
                f' onclick="toggleSeverityFilter(this)"'
                f' title="Filter: {cap_sev} (multi-select)">'
                f'<span class="chip-dot" style="background:{color};"></span>'
                f'<span class="chip-count">{count}</span>'
                f'<span class="chip-label">{cap_sev}</span>'
                f'</button>'
            )
    if not severity_chips:
        severity_chips = '<span style="color:var(--text-muted);">No issues found.</span>'

    # ---- Score summary ----
    if not findings:
        score_summary = "The domain passed all checks with an excellent security posture."
    else:
        score_summary = f"{len(findings)} finding(s) identified. Click a severity chip to filter."

    # ---- Category sub-scores HTML ----
    cat_scores_html = _category_scores_html(category_scores or {})

    # CSS (plain string — no f-string, no {{ }} escaping needed)
    # ================================================================
    css = """
:root {
  --bg:#f8fafc; --card-bg:#ffffff;
  --card-shadow:0 1px 3px rgba(0,0,0,0.08),0 1px 2px rgba(0,0,0,0.04);
  --text-primary:#1e293b; --text-secondary:#475569; --text-muted:#94a3b8;
  --border:#e2e8f0; --header-bg:#1e293b; --header-text:#f8fafc;
  --rec-bg:#f1f5f9; --toggle-bg:#e2e8f0; --toggle-knob:#ffffff;
  --sidebar-bg:#1e293b; --sidebar-text:#cbd5e1; --sidebar-active:#3b82f6;
  --sidebar-hover:rgba(255,255,255,0.08); --sidebar-border:#334155;
  --sidebar-width:280px;
  --ref-bg:#f8fafc;
  --ref-border:#e2e8f0;
  --ref-link:#2563eb;
}
[data-theme="dark"] {
  --bg:#0f172a; --card-bg:#1e293b;
  --card-shadow:0 1px 3px rgba(0,0,0,0.4);
  --text-primary:#f1f5f9; --text-secondary:#94a3b8; --text-muted:#64748b;
  --border:#334155; --header-bg:#020617; --header-text:#f1f5f9;
  --rec-bg:#0f172a; --toggle-bg:#3b82f6; --toggle-knob:#ffffff;
  --sidebar-bg:#020617; --sidebar-text:#94a3b8; --sidebar-active:#60a5fa;
  --sidebar-hover:rgba(255,255,255,0.06); --sidebar-border:#1e293b;
  --ref-bg:#1e293b; --ref-border:#334155; --ref-link:#60a5fa;
}
*{box-sizing:border-box;margin:0;padding:0;}
body{font-family:-apple-system,BlinkMacSystemFont,'Segoe UI',Roboto,sans-serif;
  background:var(--bg);color:var(--text-primary);
  transition:background .3s,color .3s;min-height:100vh;}
.app-shell{display:flex;min-height:100vh;}
.sidebar{width:var(--sidebar-width);min-width:var(--sidebar-width);
  background:var(--sidebar-bg);color:var(--sidebar-text);
  position:sticky;top:0;height:100vh;overflow-y:auto;
  display:flex;flex-direction:column;
  border-right:1px solid var(--sidebar-border);
  transition:width .25s;z-index:90;}
.sidebar.collapsed{width:0;min-width:0;overflow:hidden;}
.main-col{flex:1;min-width:0;display:flex;flex-direction:column;}
header{background:var(--header-bg);color:var(--header-text);
  padding:16px 24px;display:flex;justify-content:space-between;
  align-items:center;flex-wrap:wrap;gap:12px;
  position:sticky;top:0;z-index:80;box-shadow:0 2px 8px rgba(0,0,0,.3);}
header h1{font-size:1.3rem;font-weight:700;letter-spacing:-.02em;}
.subtitle{font-size:.82rem;opacity:.65;margin-top:2px;}
.header-right{display:flex;align-items:center;gap:16px;}
.sb-toggle-btn{background:none;border:1px solid rgba(255,255,255,.2);
  color:var(--header-text);border-radius:6px;padding:6px 10px;
  cursor:pointer;font-size:1rem;line-height:1;}
.toggle-wrapper{display:flex;align-items:center;gap:8px;font-size:.85rem;}
.toggle-label{color:var(--header-text);opacity:.85;}
.toggle{position:relative;width:44px;height:24px;cursor:pointer;}
.toggle input{opacity:0;width:0;height:0;}
.toggle-track{position:absolute;inset:0;background:var(--toggle-bg);
  border-radius:24px;transition:background .3s;}
.toggle-track::after{content:'';position:absolute;width:18px;height:18px;
  background:var(--toggle-knob);border-radius:50%;
  top:3px;left:3px;transition:transform .3s;
  box-shadow:0 1px 3px rgba(0,0,0,.2);}
input:checked + .toggle-track::after{transform:translateX(20px);}
.sidebar-header{padding:16px 16px 12px;border-bottom:1px solid var(--sidebar-border);
  font-size:.7rem;font-weight:700;letter-spacing:.1em;
  text-transform:uppercase;opacity:.6;}
.sidebar-search{margin:10px 12px;position:relative;}
.sidebar-search input{width:100%;background:rgba(255,255,255,.07);
  border:1px solid var(--sidebar-border);border-radius:6px;
  padding:6px 10px;color:var(--sidebar-text);
  font-size:.8rem;outline:none;}
.sidebar-search input::placeholder{opacity:.4;}
.sidebar-nav{flex:1;overflow-y:auto;padding:8px 0;}
.cat-group{margin-bottom:2px;}
.cat-toggle{width:100%;background:none;border:none;cursor:pointer;
  display:flex;align-items:center;gap:8px;padding:8px 16px;
  color:var(--sidebar-text);font-size:.82rem;font-weight:600;
  text-align:left;transition:background .15s;}
.cat-toggle:hover{background:var(--sidebar-hover);}
.cat-arrow{font-size:.65rem;transition:transform .2s;flex-shrink:0;opacity:.7;}
.cat-toggle[aria-expanded="false"] .cat-arrow{transform:rotate(-90deg);}
.cat-name{flex:1;text-overflow:ellipsis;overflow:hidden;white-space:nowrap;}
.cat-count{background:rgba(255,255,255,.12);border-radius:10px;
  padding:1px 7px;font-size:.72rem;flex-shrink:0;}
.cat-items{overflow:hidden;transition:max-height .25s ease;}
.cat-items.collapsed{max-height:0;}
.sidebar-item{display:flex;align-items:center;gap:8px;padding:6px 16px 6px 24px;
  font-size:.78rem;color:var(--sidebar-text);text-decoration:none;
  transition:background .12s;cursor:pointer;
  border-left:2px solid transparent;}
.sidebar-item:hover{background:var(--sidebar-hover);}
.sidebar-item.active{background:rgba(59,130,246,.15);
  border-left-color:var(--sidebar-active);color:#fff;}
.sev-dot{width:7px;height:7px;border-radius:50%;flex-shrink:0;}
.sidebar-item-text{overflow:hidden;text-overflow:ellipsis;white-space:nowrap;}
.sidebar-empty{padding:16px;font-size:.82rem;opacity:.5;}
.sidebar-footer{padding:10px 16px;border-top:1px solid var(--sidebar-border);
  font-size:.7rem;opacity:.4;}
.container{max-width:960px;margin:0 auto;padding:28px 20px;}
.score-section{display:flex;align-items:center;gap:36px;
  background:var(--card-bg);border-radius:12px;
  padding:28px 36px;margin-bottom:24px;
  box-shadow:var(--card-shadow);flex-wrap:wrap;}
.gauge-wrap{position:relative;width:180px;height:180px;flex-shrink:0;}
.gauge-wrap svg{transform:rotate(-90deg);}
.gauge-center{position:absolute;inset:0;display:flex;flex-direction:column;
  align-items:center;justify-content:center;}
.gauge-score{font-size:2.8rem;font-weight:800;line-height:1;}
.gauge-label{font-size:.78rem;color:var(--text-muted);margin-top:4px;}
.gauge-grade{font-size:1.6rem;font-weight:700;width:40px;height:40px;
  border-radius:50%;display:flex;align-items:center;
  justify-content:center;color:#fff;margin-top:6px;}
.score-info h2{font-size:1.2rem;margin-bottom:12px;}
.chips-area{display:flex;flex-wrap:wrap;gap:4px;margin-top:12px;}
.sev-chip{display:inline-flex;align-items:center;gap:6px;
  border:2px solid transparent;border-radius:20px;
  padding:6px 14px;cursor:pointer;font-size:.9rem;
  background:var(--card-bg);color:var(--text-primary);
  transition:all .18s ease;outline:none;
  box-shadow:var(--card-shadow);}
.sev-chip:hover{transform:translateY(-1px);
  box-shadow:0 4px 10px rgba(0,0,0,.15);}
.sev-chip.active{transform:translateY(-1px);}
.chip-dot{width:10px;height:10px;border-radius:50%;flex-shrink:0;}
.chip-count{font-weight:700;}
.chip-label{text-transform:capitalize;color:var(--text-muted);}
.meta-grid{display:grid;grid-template-columns:repeat(auto-fill,minmax(180px,1fr));
  gap:12px;background:var(--card-bg);border-radius:12px;
  padding:20px;margin-bottom:24px;box-shadow:var(--card-shadow);}
.meta-key{font-size:.72rem;text-transform:uppercase;letter-spacing:.08em;
  color:var(--text-muted);margin-bottom:4px;}
.meta-val{font-weight:600;font-size:.9rem;word-break:break-all;}
.section-header{display:flex;align-items:center;justify-content:space-between;
  margin-bottom:14px;flex-wrap:wrap;gap:8px;}
.section-header h2{font-size:1.1rem;}
.finding-card{transition:box-shadow .2s;}
.finding-card:hover{box-shadow:0 4px 12px rgba(0,0,0,.12)!important;}
.filter-hint{font-size:.82rem;color:var(--text-muted);font-style:italic;}
.clear-btn{background:none;border:1px solid var(--border);border-radius:12px;
  padding:3px 12px;font-size:.8rem;color:var(--text-muted);
  cursor:pointer;display:none;}
.clear-btn:hover{color:var(--text-primary);border-color:var(--text-primary);}
#no-results{display:none;text-align:center;padding:40px;
  color:var(--text-muted);font-size:.95rem;}
footer{text-align:center;padding:20px;color:var(--text-muted);font-size:.8rem;
  border-top:1px solid var(--border);margin-top:36px;}
@media(max-width:768px){
  .sidebar{position:fixed;left:0;top:0;height:100vh;
    transform:translateX(-100%);transition:transform .25s;}
  .sidebar:not(.collapsed){transform:translateX(0);}
}
/* ---- Manual Verification tabs ---- */
.verif-tabs-wrap { margin-top: 14px; }
.verif-tab-bar {
  display: flex; gap: 4px; border-bottom: 2px solid var(--border);
  margin-bottom: 14px;
}
.verif-tab {
  display: inline-flex; align-items: center; gap: 5px;
  padding: 6px 14px; border: none; border-radius: 6px 6px 0 0;
  font-size: 0.82rem; font-weight: 600; cursor: pointer;
  background: none; transition: background .15s, color .15s;
  margin-bottom: -2px; border-bottom: 2px solid transparent;
}
.verif-tab-active {
  color: #3b82f6; border-bottom-color: #3b82f6;
  background: rgba(59,130,246,0.07);
}
.verif-tab-inactive {
  color: var(--text-muted);
}
.verif-tab-inactive:hover { background: var(--rec-bg); color: var(--text-primary); }
.verif-panels { display: flex; flex-direction: column; gap: 12px; }
.verif-panel { display: flex; flex-direction: column; gap: 12px; }
.verif-card {
  border: 1px solid var(--border); border-radius: 10px;
  padding: 16px; background: var(--card-bg);
  box-shadow: 0 1px 3px rgba(0,0,0,0.06);
}
.verif-card-header { display: flex; align-items: center; gap: 10px; margin-bottom: 10px; }
.verif-icon {
  width: 32px; height: 32px; border-radius: 7px;
  display: flex; align-items: center; justify-content: center; flex-shrink: 0;
}
.verif-tool-name { font-weight: 700; font-size: 0.95rem; }
.verif-desc { font-size: 0.87rem; color: var(--text-secondary); margin: 0 0 10px; line-height: 1.55; }
.verif-steps { font-size: 0.87rem; color: var(--text-secondary); margin: 0 0 10px; padding-left: 18px; line-height: 1.6; }
.verif-code {
  background: #f5f5f5; border-radius: 6px; padding: 10px 12px;
  font-size: 0.82rem; font-family: ui-monospace,'Cascadia Code',Menlo,monospace;
  overflow-x: auto; white-space: pre; margin: 0 0 8px; color: #1e293b;
}
[data-theme="dark"] .verif-code { background: #0f172a; color: #e2e8f0; }
.verif-confirm { font-size: 0.83rem; color: var(--text-secondary); margin: 6px 0 0; line-height: 1.5; }
/* ---- Remediation ---- */
.remed-section { margin-top: 20px; }
.remed-header {
  font-size: 0.7rem; font-weight: 700; letter-spacing: 0.12em;
  text-transform: uppercase; color: var(--text-muted); margin-bottom: 12px;
}
.remed-box {
  border: 1.5px solid #86efac; border-radius: 10px;
  padding: 18px 20px; background: rgba(240,253,244,0.5);
}
[data-theme="dark"] .remed-box { background: rgba(20,83,45,0.15); border-color: #166534; }
.remed-title { font-weight: 700; font-size: 0.95rem; color: #16a34a; margin-bottom: 16px; }
[data-theme="dark"] .remed-title { color: #4ade80; }
.remed-step { display: flex; align-items: flex-start; gap: 14px; margin-bottom: 14px; }
.remed-num {
  width: 26px; height: 26px; border-radius: 50%;
  border: 2px solid #86efac; color: #16a34a; font-weight: 700;
  font-size: 0.8rem; display: flex; align-items: center;
  justify-content: center; flex-shrink: 0; margin-top: 1px;
}
[data-theme="dark"] .remed-num { border-color: #166534; color: #4ade80; }
.remed-step-body { flex: 1; }
.remed-step-body p { font-size: 0.88rem; color: var(--text-secondary); line-height: 1.55; margin: 0 0 8px; }
.remed-code { margin-top: 6px; }
/* ---- Section Toggle (details/summary) ---- */
details > summary { list-style: none; }
details > summary::-webkit-details-marker { display: none; }
details[open] > summary > span:first-child { transform: rotate(0deg) !important; }
details:not([open]) > summary > span:first-child { transform: rotate(-90deg) !important; }

/* ---- Executive Summary card ---- */
.exec-summary {
  background: var(--card-bg);
  border-radius: 12px;
  padding: 28px 32px;
  margin-bottom: 24px;
  box-shadow: var(--card-shadow);
  border-left: 5px solid var(--header-bg);
  position: relative;
  overflow: hidden;
}
.exec-summary::before {
  content: '';
  position: absolute;
  top: 0; right: 0;
  width: 200px; height: 100%;
  background: linear-gradient(135deg, transparent 55%, rgba(59,130,246,0.04) 100%);
  pointer-events: none;
}
.exec-header {
  display: flex;
  align-items: center;
  justify-content: space-between;
  margin-bottom: 18px;
  flex-wrap: wrap;
  gap: 12px;
}
.exec-title {
  font-size: 0.68rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--text-muted);
  display: flex;
  align-items: center;
  gap: 7px;
}
.exec-grade-pill {
  display: inline-flex;
  align-items: center;
  gap: 10px;
  border: 1.5px solid #86efac;
  background: #f0fdf4;
  border-radius: 24px;
  padding: 6px 16px 6px 10px;
}
.exec-grade-circle {
  width: 34px; height: 34px;
  border-radius: 50%;
  color: #fff;
  font-weight: 800;
  font-size: 1rem;
  display: flex; align-items: center; justify-content: center;
  flex-shrink: 0;
}
.exec-grade-score {
  font-size: 0.9rem;
  font-weight: 700;
  line-height: 1.1;
}
.exec-grade-label {
  font-size: 0.7rem;
  color: var(--text-muted);
}
.exec-body {
  font-size: 0.95rem;
  line-height: 1.75;
  color: var(--text-secondary);
  margin-bottom: 4px;
  width: 100%;
}
.exec-body strong { color: var(--text-primary); }
.exec-divider {
  border: none;
  border-top: 1px solid var(--border);
  margin: 20px 0 16px;
}
.exec-top-label {
  font-size: 0.68rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--text-muted);
  margin-bottom: 12px;
}
.exec-chips-row {
  display: flex;
  gap: 12px;
  flex-wrap: wrap;
}
.exec-chip {
  display: flex;
  align-items: flex-start;
  gap: 10px;
  border-radius: 8px;
  padding: 10px 14px;
  flex: 1;
  min-width: 200px;
  max-width: 320px;
  text-decoration: none;
  cursor: pointer;
  transition: transform .15s ease, box-shadow .15s ease;
}
.exec-chip:hover, .exec-chip:focus {
  transform: translateY(-2px);
  box-shadow: 0 4px 12px rgba(0,0,0,0.12);
  outline: none;
}
.exec-chip-bar {
  width: 3px;
  border-radius: 2px;
  align-self: stretch;
  flex-shrink: 0;
  min-height: 40px;
}
.exec-chip-sev {
  font-size: 0.62rem;
  font-weight: 700;
  text-transform: uppercase;
  letter-spacing: 0.08em;
  margin-bottom: 2px;
}
.exec-chip-title {
  font-size: 0.83rem;
  font-weight: 600;
  color: var(--text-primary);
  line-height: 1.35;
}
.exec-chip-meta {
  font-size: 0.72rem;
  color: var(--text-muted);
  margin-top: 3px;
}
@media print {
  .exec-summary {
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    margin-bottom: 8mm !important;
    border: 1px solid #cbd5e1 !important;
    background: #fff !important;
    box-shadow: none !important;
  }
  .exec-summary::before { display: none !important; }
  .exec-grade-pill { border: 1px solid #86efac !important; }
  .exec-chip {
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    border: 1px solid #e2e8f0 !important;
  }
  .exec-body { font-size: 0.88rem !important; }
}

/* ==================================================================
   PRINT / PDF STYLESHEET
   Optimised for A4/Letter -- use browser Print to Save as PDF
   ================================================================== */
@media print {
  /* ---- Page setup ---- */
  @page {
    size: A4;
    margin: 18mm 16mm 18mm 16mm;
  }
  @page :first {
    margin-top: 12mm;
  }

  /* ---- Force colour preservation ---- */
  * {
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
    color-adjust: exact !important;
  }

  /* ---- Hide interactive / screen-only elements ---- */
  .sidebar,
  .sb-toggle-btn,
  .toggle-wrapper,
  .chips-area,
  .sev-chip,
  .clear-btn,
  #clear-btn,
  .filter-hint,
  #filter-hint,
  .sidebar-search,
  .sidebar-footer,
  .verif-tab-bar,
  footer a[href] {
    display: none !important;
  }

  /* ---- Collapse the app shell to a single column ---- */
  body, html {
    background: #fff !important;
    color: #1a1a1a !important;
  }

  .app-shell {
    display: block !important;
  }

  .main-col {
    display: block !important;
    width: 100% !important;
  }

  /* ---- Header ---- */
  header {
    position: static !important;
    background: #1e293b !important;
    color: #fff !important;
    padding: 12px 16px !important;
    box-shadow: none !important;
    margin-bottom: 8mm !important;
    border-radius: 0 !important;
  }

  header h1  { font-size: 1.2rem !important; }
  .subtitle  { font-size: 0.75rem !important; }

  /* ---- Container ---- */
  .container {
    max-width: 100% !important;
    padding: 0 !important;
    margin: 0 !important;
  }

  /* ---- Score section ---- */
  .score-section {
    display: flex !important;
    align-items: center !important;
    gap: 24px !important;
    padding: 16px 20px !important;
    margin-bottom: 8mm !important;
    border: 1px solid #cbd5e1 !important;
    box-shadow: none !important;
    border-radius: 8px !important;
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    background: #fff !important;
  }

  .score-info h2 { font-size: 1rem !important; }

  /* ---- Metadata grid ---- */
  .meta-grid {
    display: grid !important;
    grid-template-columns: repeat(3, 1fr) !important;
    gap: 8px !important;
    padding: 14px 16px !important;
    margin-bottom: 8mm !important;
    border: 1px solid #cbd5e1 !important;
    box-shadow: none !important;
    border-radius: 6px !important;
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    background: #fff !important;
  }

  .meta-key { font-size: 0.68rem !important; }
  .meta-val { font-size: 0.85rem !important; }

  /* ---- Section header ---- */
  .section-header {
    page-break-after: avoid !important;
    break-after: avoid !important;
    margin-bottom: 6px !important;
  }

  /* ---- Finding cards ---- */
  .finding-card {
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    margin-bottom: 8mm !important;
    box-shadow: none !important;
    border-radius: 6px !important;
    background: #fff !important;
    color: #1a1a1a !important;
  }

  .finding-card h3 { font-size: 0.95rem !important; }

  .finding-card p,
  .finding-card li {
    color: #333 !important;
    font-size: 0.85rem !important;
  }

  /* ---- Force-open all details/summary sections ---- */
  details > *:not(summary) {
    display: block !important;
  }

  details > summary {
    page-break-after: avoid !important;
    break-after: avoid !important;
  }

  /* Suppress the animated chevron arrows */
  details > summary > span:first-child {
    display: none !important;
  }

  /* ---- Affected objects: cap at 20 entries for print ---- */
  details ul li:nth-child(n+21) {
    display: none !important;
  }

  /* ---- Verification tool cards ---- */
  .verif-tabs-wrap {
    margin-top: 8px !important;
  }

  /* Show ALL tab panels in print (not just the active one) */
  .verif-panel {
    display: block !important;
  }

  .verif-card {
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    margin-bottom: 6px !important;
    border: 1px solid #cbd5e1 !important;
    box-shadow: none !important;
    background: #fff !important;
  }

  .verif-tool-name { font-size: 0.88rem !important; }
  .verif-desc      { font-size: 0.82rem !important; }
  .verif-confirm   { font-size: 0.78rem !important; }

  /* ---- Code blocks ---- */
  .verif-code,
  .remed-code,
  pre,
  code {
    background: #f4f4f4 !important;
    color: #1a1a1a !important;
    font-size: 0.75rem !important;
    white-space: pre-wrap !important;
    word-break: break-all !important;
    border: 1px solid #e2e8f0 !important;
    border-radius: 4px !important;
    padding: 8px 10px !important;
    page-break-inside: avoid !important;
    break-inside: avoid !important;
  }

  /* ---- Remediation box ---- */
  .remed-box {
    background: #f0fdf4 !important;
    border: 1.5px solid #86efac !important;
    page-break-inside: avoid !important;
    break-inside: avoid !important;
  }

  .remed-title { color: #16a34a !important; }
  .remed-num   { color: #16a34a !important; border-color: #86efac !important; }

  .remed-step {
    page-break-inside: avoid !important;
    break-inside: avoid !important;
  }

  .remed-step-body p { font-size: 0.83rem !important; }

  /* ---- References: print URLs inline ---- */
  details a[href]::after {
    content: " (" attr(href) ")";
    font-size: 0.7rem;
    color: #555;
    word-break: break-all;
  }

  header a::after,
  footer a::after {
    content: none !important;
  }

  /* ---- Footer ---- */
  footer {
    display: block !important;
    text-align: center !important;
    font-size: 0.72rem !important;
    color: #555 !important;
    border-top: 1px solid #e2e8f0 !important;
    margin-top: 8mm !important;
    padding-top: 4mm !important;
    page-break-before: avoid !important;
  }
}

/* ---- Category score grid ---- */
.cat-grid-section {
  background: var(--card-bg);
  border-radius: 12px;
  padding: 20px 24px;
  margin-bottom: 24px;
  box-shadow: var(--card-shadow);
}
.cat-grid-label {
  font-size: 0.68rem;
  font-weight: 700;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--text-muted);
  margin-bottom: 14px;
}
.cat-grid {
  display: grid;
  grid-template-columns: repeat(auto-fill, minmax(200px, 1fr));
  gap: 12px;
}
.cat-card {
  background: var(--rec-bg);
  border-radius: 8px;
  padding: 12px 14px;
  border: 1px solid var(--border);
}
.cat-card:hover {
  border-color: var(--sidebar-active);
  box-shadow: 0 2px 8px rgba(59,130,246,0.15);
  transform: translateY(-1px);
  transition: all .15s ease;
}
.cat-card.cat-active {
  border-color: var(--sidebar-active);
  box-shadow: 0 0 0 2px var(--sidebar-active);
  background: rgba(59,130,246,0.07);
}

.cat-card-header {
  display: flex;
  justify-content: space-between;
  align-items: center;
  margin-bottom: 8px;
}
.cat-card-name {
  font-size: 0.82rem;
  font-weight: 600;
  color: var(--text-primary);
  overflow: hidden;
  text-overflow: ellipsis;
  white-space: nowrap;
  max-width: 75%;
}
.cat-card-grade {
  font-size: 0.72rem;
  font-weight: 700;
  color: #fff;
  border-radius: 4px;
  padding: 1px 7px;
  flex-shrink: 0;
}
.cat-bar-track {
  height: 6px;
  background: var(--border);
  border-radius: 3px;
  overflow: hidden;
  margin-bottom: 6px;
}
.cat-bar-fill {
  height: 100%;
  border-radius: 3px;
  transition: width 0.5s ease;
}
.cat-card-score {
  font-size: 0.75rem;
  font-weight: 600;
  text-align: right;
}
@media print {
  .cat-grid-section {
    page-break-inside: avoid !important;
    break-inside: avoid !important;
    border: 1px solid #cbd5e1 !important;
    box-shadow: none !important;
    background: #fff !important;
  }
  .cat-card {
    border: 1px solid #e2e8f0 !important;
    background: #f8fafc !important;
  }
  .cat-bar-track {
    background: #e2e8f0 !important;
  }
}
""".strip()

    # ================================================================
    # JavaScript (plain string — no f-string, no {{ }} escaping needed)
    # ================================================================
    js = """
// ---- Theme ----
(function() {
  var saved = localStorage.getItem('adscan-theme');
  if (saved === 'dark') {
    document.documentElement.setAttribute('data-theme', 'dark');
    var t = document.getElementById('darkToggle');
    if (t) t.checked = true;
  }
})();
function toggleTheme(el) {
  var d = el.checked;
  document.documentElement.setAttribute('data-theme', d ? 'dark' : 'light');
  localStorage.setItem('adscan-theme', d ? 'dark' : 'light');
}
// ---- Sidebar ----
function toggleSidebar() {
  document.getElementById('sidebar').classList.toggle('collapsed');
}
// ---- Category collapse ----
function toggleCat(btn) {
  var expanded = btn.getAttribute('aria-expanded') === 'true';
  btn.setAttribute('aria-expanded', expanded ? 'false' : 'true');
  var items = btn.nextElementSibling;
  if (expanded) {
    items.style.maxHeight = items.scrollHeight + 'px';
    items.style.overflow = 'hidden';
    requestAnimationFrame(function() {
      requestAnimationFrame(function() { items.style.maxHeight = '0'; });
    });
    items.addEventListener('transitionend', function h() {
      items.classList.add('collapsed'); items.style.maxHeight = '';
      items.removeEventListener('transitionend', h);
    });
  } else {
    items.classList.remove('collapsed');
    items.style.overflow = 'hidden';
    items.style.maxHeight = items.scrollHeight + 'px';
    items.addEventListener('transitionend', function h() {
      items.style.maxHeight = 'none'; items.style.overflow = '';
      items.removeEventListener('transitionend', h);
    });
  }
}
// ---- Sidebar search ----
function sidebarSearchFilter(val) {
  var q = val.toLowerCase().trim();
  document.querySelectorAll('.cat-group').forEach(function(grp) {
    var any = false;
    grp.querySelectorAll('.sidebar-item').forEach(function(item) {
      var match = !q || item.querySelector('.sidebar-item-text').textContent.toLowerCase().includes(q);
      item.style.display = match ? '' : 'none';
      if (match) any = true;
    });
    grp.style.display = any ? '' : 'none';
  });
}
// ---- Sidebar active ----
function sidebarNav(link, e) {
  document.querySelectorAll('.sidebar-item').forEach(function(a) { a.classList.remove('active'); });
  link.classList.add('active');
}
// ---- Intersection observer ----
(function() {
  var cards = document.querySelectorAll('.finding-card[id]');
  if (!cards.length) return;
  var obs = new IntersectionObserver(function(entries) {
    entries.forEach(function(entry) {
      if (!entry.isIntersecting) return;
      var id = entry.target.id;
      document.querySelectorAll('.sidebar-item').forEach(function(a) {
        var match = a.getAttribute('href') === '#' + id;
        a.classList.toggle('active', match);
        if (match) a.scrollIntoView({block:'nearest',behavior:'smooth'});
      });
    });
  }, {threshold: 0.35});
  cards.forEach(function(c) { obs.observe(c); });
})();
// ---- Severity chip filter ----
var activeFilters = {};
var activeCatFilters = {};
function toggleSeverityFilter(btn) {
  var sev = btn.getAttribute('data-sev');
  var color = btn.getAttribute('data-color') || '#6b7280';
  if (activeFilters[sev]) {
    delete activeFilters[sev];
    btn.classList.remove('active');
    btn.style.borderColor = 'transparent';
    btn.style.color = '';
  } else {
    activeFilters[sev] = true;
    btn.classList.add('active');
    btn.style.borderColor = color;
    btn.style.color = color;
  }
  applyFilters();
}
function toggleCategoryFilter(card) {
  var slug = card.getAttribute('data-catslug');
  if (!slug) return;
  if (activeCatFilters[slug]) {
    delete activeCatFilters[slug];
    card.classList.remove('cat-active');
  } else {
    activeCatFilters[slug] = true;
    card.classList.add('cat-active');
  }
  applyFilters();
}

function applyFilters() {
  var sevKeys = Object.keys(activeFilters);
  var catKeys = Object.keys(activeCatFilters);
  var cards = document.querySelectorAll('.finding-card');
  var clearBtn = document.getElementById('clear-btn');
  var hintEl = document.getElementById('filter-hint');
  var countEl = document.getElementById('visible-count');
  var noRes = document.getElementById('no-results');
  var vis = 0;
  var hasFilter = sevKeys.length > 0 || catKeys.length > 0;
  if (!hasFilter) {
    cards.forEach(function(c) { c.style.display = ''; });
    vis = cards.length;
    if (clearBtn) clearBtn.style.display = 'none';
    if (hintEl) hintEl.textContent = '';
  } else {
    cards.forEach(function(c) {
      var sevMatch = sevKeys.length === 0 || !!activeFilters[c.getAttribute('data-severity')];
      var cardCats = (c.getAttribute('data-category') || '').split(' ');
      var catMatch = catKeys.length === 0 || cardCats.some(function(slug) { return !!activeCatFilters[slug]; });
      var show = sevMatch && catMatch;
      c.style.display = show ? '' : 'none';
      if (show) vis++;
    });
    if (clearBtn) clearBtn.style.display = 'inline-block';
    var hints = [];
    if (sevKeys.length) hints.push('severity: ' + sevKeys.join(', '));
    if (catKeys.length) hints.push('category: ' + catKeys.join(', '));
    if (hintEl) hintEl.textContent = 'Showing: ' + hints.join(' | ');
  }
  if (countEl) countEl.textContent = vis;
  if (noRes) noRes.style.display = vis === 0 ? 'block' : 'none';
}
function clearFilters() {
  activeFilters = {};
  activeCatFilters = {};
  document.querySelectorAll('.sev-chip').forEach(function(b) {
    b.classList.remove('active');
    b.style.borderColor = 'transparent';
    b.style.color = '';
  });
  document.querySelectorAll('.cat-card').forEach(function(c) {
    c.classList.remove('cat-active');
  });
  applyFilters();
}

// ---- Verification tab switcher ----
function verifTab(btn, panelId) {
  var bar = btn.parentElement;
  bar.querySelectorAll('.verif-tab').forEach(function(b) {
    b.classList.remove('verif-tab-active');
    b.classList.add('verif-tab-inactive');
  });
  btn.classList.remove('verif-tab-inactive');
  btn.classList.add('verif-tab-active');
  var wrap = bar.nextElementSibling; // .verif-panels
  wrap.querySelectorAll('.verif-panel').forEach(function(p) {
    p.style.display = 'none';
  });
  var target = document.getElementById(panelId);
  if (target) target.style.display = '';
}
""".strip()

    # ================================================================
    # HTML (f-string only for Python values — NO JS or CSS braces)
    # ================================================================
    html_content = f"""<!DOCTYPE html>
<html lang="en" data-theme="light">
<head>
  <meta charset="UTF-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>ADScan Report - {html_mod.escape(domain)}</title>
  <style>{css}</style>
</head>
<body>
<div class="app-shell">
  <aside class="sidebar" id="sidebar">
    <div class="sidebar-header">Category Navigator</div>
    <div class="sidebar-search">
      <input type="text" id="sidebarSearch" placeholder="&#128269; Filter checks..."
             oninput="sidebarSearchFilter(this.value)">
    </div>
    <nav class="sidebar-nav" id="sidebarNav">
      {sidebar_items_html}
    </nav>
    <div class="sidebar-footer">ADScan &bull; {html_mod.escape(domain)}</div>
  </aside>
  <div class="main-col">
    <header>
      <div style="display:flex;align-items:center;gap:12px;">
        <button class="sb-toggle-btn" onclick="toggleSidebar()" title="Toggle sidebar">&#9776;</button>
        <div>
          <h1>&#x1F6E1; ADScan Report</h1>
          <div class="subtitle">Active Directory Vulnerability Assessment &mdash; {html_mod.escape(domain)}</div>
        </div>
      </div>
      <div class="header-right">
        <div class="toggle-wrapper">
          <span class="toggle-label">&#9788;</span>
          <label class="toggle">
            <input type="checkbox" id="darkToggle" onchange="toggleTheme(this)">
            <span class="toggle-track"></span>
          </label>
          <span class="toggle-label">&#9790;</span>
        </div>
      </div>
    </header>
    {_exec_summary_html(domain, dc_host, scan_time, score, grade, score_color, sev_counts, findings)}
    <div class="container">
      <div class="score-section">
        <div class="gauge-wrap">
          <svg width="180" height="180" viewBox="0 0 200 200">
            <circle cx="100" cy="100" r="{radius}" fill="none" stroke="var(--border)" stroke-width="16"/>
            <circle cx="100" cy="100" r="{radius}" fill="none" stroke="{score_color}" stroke-width="16"
              stroke-dasharray="{circumference:.2f}" stroke-dashoffset="{dash_offset:.2f}"
              stroke-linecap="round"/>
          </svg>
          <div class="gauge-center">
            <div class="gauge-score" style="color:{score_color};">{score}</div>
            <div class="gauge-label">/ 100</div>
            <div class="gauge-grade" style="background:{score_color};">{grade}</div>
          </div>
        </div>
        <div class="score-info">
          <h2>Security Score</h2>
          <p style="color:var(--text-secondary);margin-bottom:14px;line-height:1.6;">{score_summary}</p>
          <div class="chips-area">{severity_chips}</div>
        </div>
      </div>
      <div class="meta-grid">
        <div><div class="meta-key">Domain</div><div class="meta-val">{html_mod.escape(domain)}</div></div>
        <div><div class="meta-key">Domain Controller</div><div class="meta-val">{html_mod.escape(dc_host)}</div></div>
        <div><div class="meta-key">Username</div><div class="meta-val">{html_mod.escape(username)}</div></div>
        <div><div class="meta-key">Protocol(s)</div><div class="meta-val">{html_mod.escape(', '.join(p.upper() for p in protocols))}</div></div>
        <div><div class="meta-key">Scan Time</div><div class="meta-val">{html_mod.escape(scan_time)}</div></div>
        <div><div class="meta-key">Total Findings</div><div class="meta-val">{len(findings)}</div></div>
      </div>
      {cat_scores_html}
      <div class="section-header">
        <h2>Findings (<span id="visible-count">{len(findings)}</span>)</h2>
        <div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap;">
          <span class="filter-hint" id="filter-hint"></span>
          <button class="clear-btn" id="clear-btn" onclick="clearFilters()">&times; Clear filter</button>
        </div>
      </div>
      <div id="no-results">No findings match the selected filter(s).
        <a href="#" onclick="clearFilters();return false;" style="color:inherit;">Clear</a>
      </div>
      {cards_html}
    </div>
    <footer>
      Generated by <strong>ADScan</strong> &mdash; {html_mod.escape(scan_time)} &mdash;
      <a href="https://github.com/dehobbs/ADScan"
         style="color:inherit;text-decoration:underline;" target="_blank">
        github.com/dehobbs/ADScan
      </a>
    </footer>
  </div>
</div>
<script>{js}</script>
</body>
</html>"""

    with open(output_file, "w", encoding="utf-8") as fh:
        fh.write(html_content)


# ---------------------------------------------------------------------------
# JSON output
# ---------------------------------------------------------------------------
def generate_json_report(output_file, domain, dc_host, username, protocols, findings, score, category_scores=None):
    """Write a machine-readable JSON report."""
    import json
    from datetime import datetime

    def _grade(s):
        if s >= 90: return "A"
        if s >= 75: return "B"
        if s >= 60: return "C"
        if s >= 40: return "D"
        return "F"

    payload = {
        "meta": {
            "tool": "ADScan",
            "version": "1.1",
            "scan_time": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
            "domain": domain,
            "dc_host": dc_host,
            "username": username,
            "protocols": protocols,
        },
        "score": {
            "value": score,
            "grade": _grade(score),
        },
        "summary": {
            "total_findings": len(findings),
            "critical": sum(1 for f in findings if f.get("severity", "").lower() == "critical"),
            "high":     sum(1 for f in findings if f.get("severity", "").lower() == "high"),
            "medium":   sum(1 for f in findings if f.get("severity", "").lower() == "medium"),
            "low":      sum(1 for f in findings if f.get("severity", "").lower() == "low"),
            "info":     sum(1 for f in findings if f.get("severity", "").lower() == "info"),
        },
        "findings": [
            {
                "title":          f.get("title", ""),
                "severity":       f.get("severity", "info"),
                "category":       f.get("category", "Uncategorized"),
                "deduction":      f.get("deduction", 0),
                "description":    f.get("description", ""),
                "recommendation": f.get("recommendation", ""),
                "affected_count": f.get("affected_count", len(f.get("details", []))),
                "details":        f.get("details", []),
            }
            for f in findings
        ],
    }

    with open(output_file, "w", encoding="utf-8") as fh:
        json.dump(payload, fh, indent=2, default=str)


# ---------------------------------------------------------------------------
# CSV output
# ---------------------------------------------------------------------------
def generate_csv_report(output_file, domain, dc_host, username, protocols, findings, score, category_scores=None):
    """Write a flat CSV report — one row per finding."""
    import csv
    from datetime import datetime

    fieldnames = [
        "scan_time", "domain", "dc_host", "username", "protocols", "score",
        "title", "severity", "category", "deduction",
        "description", "recommendation", "affected_count", "details_sample",
    ]
    scan_time = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
    proto_str = ", ".join(protocols)

    with open(output_file, "w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        writer.writeheader()

        if not findings:
            writer.writerow({
                "scan_time": scan_time, "domain": domain, "dc_host": dc_host,
                "username": username, "protocols": proto_str, "score": score,
                "title": "No findings", "severity": "", "category": "",
                "deduction": 0, "description": "", "recommendation": "",
                "affected_count": 0, "details_sample": "",
            })
            return

        for f in findings:
            details = f.get("details", [])
            sample = " | ".join(str(d) for d in details[:5])
            if len(details) > 5:
                sample += f" ... (+{len(details) - 5} more)"
            cats = f.get("category", "Uncategorized")
            if not isinstance(cats, str):
                cats = ", ".join(cats)
            writer.writerow({
                "scan_time":      scan_time,
                "domain":         domain,
                "dc_host":        dc_host,
                "username":       username,
                "protocols":      proto_str,
                "score":          score,
                "title":          f.get("title", ""),
                "severity":       f.get("severity", "info"),
                "category":       cats,
                "deduction":      f.get("deduction", 0),
                "description":    f.get("description", ""),
                "recommendation": f.get("recommendation", ""),
                "affected_count": f.get("affected_count", len(details)),
                "details_sample": sample,
            })


# ---------------------------------------------------------------------------
# DOCX output
# ---------------------------------------------------------------------------
def generate_docx_report(output_file, domain, dc_host, username, protocols, findings, score, category_scores=None):
    """Write a Microsoft Word (.docx) report using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX export. "
            "Install it with: pip install python-docx"
        )
    from datetime import datetime

    def _grade(s):
        if s >= 90: return "A"
        if s >= 75: return "B"
        if s >= 60: return "C"
        if s >= 40: return "D"
        return "F"

    SEV_RGB = {
        "critical": RGBColor(0xDC, 0x26, 0x26),
        "high":     RGBColor(0xEA, 0x58, 0x0C),
        "medium":   RGBColor(0xD9, 0x77, 0x06),
        "low":      RGBColor(0x25, 0x63, 0xEB),
        "info":     RGBColor(0x6B, 0x72, 0x80),
    }
    SEV_ORDER_D = ["critical", "high", "medium", "low", "info"]

    def _set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def _set_para_bg(para, hex_color):
        """Shade a paragraph background (for code/evidence blocks)."""
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    scan_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    grade = _grade(score)
    proto_str = ", ".join(protocols) if protocols else "N/A"

    sev_counts = {s: 0 for s in SEV_ORDER_D}
    for f in findings:
        sev = f.get("severity", "info").lower()
        if sev in sev_counts:
            sev_counts[sev] += 1

    sorted_findings = sorted(
        findings,
        key=lambda f: SEV_ORDER_D.index(f.get("severity", "info").lower())
        if f.get("severity", "info").lower() in SEV_ORDER_D else 99,
    )

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # --- Cover block ---
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("ADScan Security Report")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    doc.add_paragraph()

    meta_tbl = doc.add_table(rows=5, cols=2)
    meta_tbl.style = "Table Grid"
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate([
        ("Domain",    domain or "N/A"),
        ("DC Host",   dc_host or "N/A"),
        ("Username",  username or "N/A"),
        ("Protocols", proto_str),
        ("Scan Time", scan_time),
    ]):
        lc = meta_tbl.rows[i].cells[0]
        vc = meta_tbl.rows[i].cells[1]
        _set_cell_bg(lc, "EFF6FF")
        _lc_run = lc.paragraphs[0].add_run(label)
        _lc_run.bold = True
        _lc_run.font.size = Pt(10)
        _vc_run = vc.paragraphs[0].add_run(value)
        _vc_run.font.size = Pt(10)

    doc.add_paragraph()
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = score_para.add_run(f"Security Score: {score}/100  \u2014  Grade: {grade}")
    sr.bold = True
    sr.font.size = Pt(16)
    sr.font.color.rgb = (
        RGBColor(0x16, 0xA3, 0x4A) if score >= 75
        else RGBColor(0xD9, 0x77, 0x06) if score >= 50
        else RGBColor(0xDC, 0x26, 0x26)
    )
    doc.add_page_break()

    # --- Executive Summary ---
    h1 = doc.add_heading("Executive Summary", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    sum_tbl = doc.add_table(rows=2, cols=6)
    sum_tbl.style = "Table Grid"
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_labels = ["Critical", "High", "Medium", "Low", "Info", "Total"]
    hdr_bgs    = ["FEF2F2",  "FFF7ED", "FFFBEB", "EFF6FF", "F9FAFB", "F3F4F6"]
    for j, (hdr, bg) in enumerate(zip(hdr_labels, hdr_bgs)):
        cell = sum_tbl.rows[0].cells[j]
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        r.bold = True
        r.font.size = Pt(9)
        if hdr.lower() in SEV_RGB:
            r.font.color.rgb = SEV_RGB[hdr.lower()]
    for j, val in enumerate([sev_counts.get(s, 0) for s in SEV_ORDER_D] + [len(findings)]):
        cell = sum_tbl.rows[1].cells[j]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(val))
        r.bold = True
        r.font.size = Pt(11)
    doc.add_paragraph()

    top = [f for f in sorted_findings if f.get("severity", "").lower() in ("critical", "high")][:10]
    if top:
        doc.add_heading("Top Priority Findings", level=2)
        for f in top:
            sev = f.get("severity", "info").lower()
            p = doc.add_paragraph(style="List Bullet")
            sr2 = p.add_run(f"[{sev.upper()}] ")
            sr2.bold = True
            sr2.font.color.rgb = SEV_RGB.get(sev, RGBColor(0, 0, 0))
            p.add_run(f.get("title", "(untitled)"))
    doc.add_page_break()

    # --- Findings Overview ---
    doc.add_heading("Findings Overview", level=1).runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    if not findings:
        doc.add_paragraph("No findings were identified during this scan.")
    else:
        ov_tbl = doc.add_table(rows=1, cols=4)
        ov_tbl.style = "Table Grid"
        for j, h in enumerate(["Severity", "Title", "Category", "Deduction"]):
            cell = ov_tbl.rows[0].cells[j]
            _set_cell_bg(cell, "1E40AF")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(9)
        sev_bg = {"critical":"FEF2F2","high":"FFF7ED","medium":"FFFBEB","low":"EFF6FF","info":"F9FAFB"}
        for f in sorted_findings:
            sev = f.get("severity", "info").lower()
            cats = f.get("category", "Uncategorized")
            if not isinstance(cats, str): cats = ", ".join(cats)
            trow = ov_tbl.add_row()
            sc = trow.cells[0]
            _set_cell_bg(sc, sev_bg.get(sev, "F9FAFB"))
            sp = sc.paragraphs[0]
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr3 = sp.add_run(sev.upper())
            sr3.bold = True
            sr3.font.size = Pt(8)
            sr3.font.color.rgb = SEV_RGB.get(sev, RGBColor(0, 0, 0))
            trow.cells[1].paragraphs[0].add_run(f.get("title", "")).font.size = Pt(9)
            trow.cells[2].paragraphs[0].add_run(cats).font.size = Pt(9)
            dp = trow.cells[3].paragraphs[0]
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ded = f.get("deduction", 0)
            dr = dp.add_run(f"-{ded}" if ded else "0")
            dr.font.size = Pt(9)
            if ded: dr.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    doc.add_page_break()

    # --- Detailed Findings ---
    doc.add_heading("Detailed Findings", level=1).runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    for idx, f in enumerate(sorted_findings, start=1):
        sev = f.get("severity", "info").lower()
        title = f.get("title", "(untitled)")
        h2 = doc.add_heading(f"{idx}. {title}", level=2)
        for r in h2.runs:
            r.font.color.rgb = SEV_RGB.get(sev, RGBColor(0, 0, 0))
        meta_p = doc.add_paragraph()
        mr = meta_p.add_run(f"Severity: {sev.upper()}")
        mr.bold = True
        mr.font.color.rgb = SEV_RGB.get(sev, RGBColor(0, 0, 0))
        cats = f.get("category", "Uncategorized")
        if not isinstance(cats, str): cats = ", ".join(cats)
        meta_p.add_run(f"  |  Category: {cats}")
        ded = f.get("deduction", 0)
        if ded: meta_p.add_run(f"  |  Score Deduction: -{ded}")
        desc = f.get("description", "")
        if desc:
            doc.add_heading("Description", level=3)
            doc.add_paragraph(desc)
        rec = f.get("recommendation", "")
        if rec:
            doc.add_heading("Recommendation", level=3)
            doc.add_paragraph(rec)
        # Raw command + output block (e.g. nxc, certipy)
        raw_output = f.get("raw_output", "")
        if raw_output:
            doc.add_heading("Command & Output", level=3)
            try:
                from PIL import Image, ImageDraw, ImageFont
                import io as _io
                _pad = 16
                _font_size = 13
                _line_spacing = 4
                _border = 2
                # Try to load a monospace font, fall back to default
                _pil_font = None
                for _fp in [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationMono-Regular.ttf",
                    "/usr/share/fonts/truetype/freefont/FreeMono.ttf",
                    "C:/Windows/Fonts/cour.ttf",
                    "C:/Windows/Fonts/consola.ttf",
                ]:
                    try:
                        _pil_font = ImageFont.truetype(_fp, _font_size)
                        break
                    except Exception:
                        pass
                if _pil_font is None:
                    _pil_font = ImageFont.load_default()
                _lines = raw_output.split("\n")
                # Measure dimensions
                _dummy_img = Image.new("RGB", (1, 1))
                _dummy_draw = ImageDraw.Draw(_dummy_img)
                _line_heights = []
                _line_widths = []
                for _ln in _lines:
                    _bb = _dummy_draw.textbbox((0, 0), _ln if _ln else " ", font=_pil_font)
                    _line_widths.append(_bb[2] - _bb[0])
                    _line_heights.append(_bb[3] - _bb[1])
                _max_w = max(_line_widths) if _line_widths else 100
                _line_h = max(_line_heights) if _line_heights else _font_size + 2
                _img_w = _max_w + _pad * 2 + _border * 2
                _img_h = ((_line_h + _line_spacing) * len(_lines)) + _pad * 2 + _border * 2
                # Draw image
                _img = Image.new("RGB", (_img_w, _img_h), color=(255, 255, 255))
                _draw = ImageDraw.Draw(_img)
                # Border
                _draw.rectangle(
                    [0, 0, _img_w - 1, _img_h - 1],
                    outline=(180, 180, 180),
                    width=_border,
                )
                # Text
                _y = _pad + _border
                for _ln in _lines:
                    _draw.text((_pad + _border, _y), _ln, font=_pil_font, fill=(30, 41, 59))
                    _y += _line_h + _line_spacing
                # Save to bytes
                _buf = _io.BytesIO()
                _img.save(_buf, format="PNG")
                _buf.seek(0)
                # Embed in DOCX — scale to fit page width (approx 16cm)
                _max_width_cm = 16.0
                _dpi = 96
                _img_w_cm = (_img_w / _dpi) * 2.54
                if _img_w_cm > _max_width_cm:
                    _scale = _max_width_cm / _img_w_cm
                    _display_w = Cm(_max_width_cm)
                else:
                    _display_w = Cm(_img_w_cm)
                _img_para = doc.add_paragraph()
                _img_para.paragraph_format.space_before = Pt(4)
                _img_para.paragraph_format.space_after = Pt(6)
                _img_run = _img_para.add_run()
                _img_run.add_picture(_buf, width=_display_w)
            except Exception as _pil_err:
                # Fallback: plain styled paragraph if Pillow unavailable
                _cmd_block_para = doc.add_paragraph()
                _cmd_block_para.paragraph_format.space_before = Pt(2)
                _cmd_block_para.paragraph_format.space_after = Pt(2)
                _cmd_block_para.paragraph_format.left_indent = Cm(0.3)
                _set_para_bg(_cmd_block_para, "F8FAFC")
                _cmd_run = _cmd_block_para.add_run(raw_output)
                _cmd_run.font.name = "Courier New"
                _cmd_run.font.size = Pt(7.5)
                _cmd_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        details = f.get("details", [])
        if details:
            n = len(details)
            doc.add_heading(f"Evidence ({n} item{'s' if n != 1 else ''})", level=3)
            # Render ALL evidence items in a monospace code-style block (no cap)
            for item in details:
                line_text = str(item).replace("[[REDACTED]]", "[REDACTED]")
                ev_p = doc.add_paragraph()
                ev_p.paragraph_format.space_before = Pt(0)
                ev_p.paragraph_format.space_after = Pt(1)
                ev_p.paragraph_format.left_indent = Cm(0.5)
                _set_para_bg(ev_p, "F1F5F9")
                ev_run = ev_p.add_run(line_text)
                ev_run.font.name = "Courier New"
                ev_run.font.size = Pt(8)
                if "[REDACTED]" in line_text:
                    ev_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
        if idx < len(sorted_findings):
            doc.add_paragraph()

    doc.save(output_file)
